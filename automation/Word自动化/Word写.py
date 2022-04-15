from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
#磅数
from docx.shared import Pt
#中文格式
from docx.oxml.ns import qn
#图片尺寸
from docx.shared import Inches
import time
#文档属性（默认字体，默认格式）——段落属性（段落间距，对齐）——文字属性（字体，自号，加粗）
price=input('请输入今日价格：')
company_list=['客户1','客户2','客户3','客户4','客户5','客户6','客户7','客户8','客户9','客户10','客户11','客户12',]
today=time.strftime('%Y-%m-%d',time.localtime())
print(today)
for i in company_list:
    document=Document()
    # 设置文档的基础字体
    document.styles['Normal'].font.name=u'楷体'
    #设置文档的基础中文字体
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'),u'楷体')
    # 初始化建立第一个自然段
    p1=document.add_paragraph()
    p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p1.add_run(f'关于下达{today}产品价格的通知')
    run1.font.name = '楷体'
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
    run1.font.size = Pt(21)
    run1.font.bold = True
    #段前段后
    p1.space_after=Pt(5)
    p1.space_before = Pt(5)

    p2 = document.add_paragraph()
    run2 = p2.add_run(i+':')
    run2.font.name = '楷体'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
    run2.font.size = Pt(16)
    run2.font.bold = True

    p3 = document.add_paragraph()
    run3 = p3.add_run(f'     根据公司安排，为提供优质客户服务，我单位拟定了今日黄金价格为{price}元，特此通知。')
    run3.font.name = '楷体'
    run3.element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
    run3.font.size = Pt(16)
    run3.font.bold = True

    p4 = document.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run4 = p4.add_run(f'联系人：小张      电话：19831958283')
    run4.font.name = '楷体'
    run4.element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
    run4.font.size = Pt(16)
    run4.font.bold = True

    table=document.add_table(rows=3,cols=3,style='Table Grid')
    table.cell(0,0).merge(table.cell(0,2))
    table_run1=table.cell(0,0).paragraphs[0].add_run('xx产品报价表')
    table_run1.font.name = '楷体'
    table_run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
    table.cell(0, 0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    #默认格式
    table.cell(1,0).text='日期'
    table.cell(1, 1).text = '价格'
    table.cell(1, 2).text = '备注'
    table.cell(2, 0).text = today
    table.cell(2, 1).text = str(price)
    table.cell(2, 2).text = ''

    document.add_page_break()#分页符

    document.add_picture('E:\图片\壁纸\女孩喝啤酒 飘窗 城市夜景 雨天 è 好看唯美4k动漫壁纸_彼岸图网.jpg',width=Inches(6))
    document.save(f'{i}-价格通知.docx')

