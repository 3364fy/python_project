from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
document=Document('E:\office\word\安全评价.docx')
document.styles['Normal'].font.name=u'楷体'
document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'),u'楷体')
document.styles['Normal'].font.size=Pt(12)
def change_text(old_text,new_text):
    all_paragraphs=document.paragraphs
    for paragraph in all_paragraphs:
        for run in paragraph:
            run_text=run.text.replace(old_text,new_text)
            run.text=run_text
    all_tables=document.tables
    for table in all_tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.replace(old_text, new_text)
                cell.text = cell_text