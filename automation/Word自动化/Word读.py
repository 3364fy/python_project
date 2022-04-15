from docx import Document
import zipfile
import re
#全是文字
document=Document('E:\office\word\安全评价.docx')
all_paragraphs=document.paragraphs
for paragraph in all_paragraphs:
    print(paragraph.text)

#包含表格
document=Document('客户1-价格通知.docx')
all_tables=document.tables
for table in all_tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)

word=zipfile.ZipFile('客户1-价格通知.docx')
xml=word.read('word/document.xml').decode('utf-8')
ex='(?<=<w:t>).*?(?=</w:t>)'
xml_list=re.findall(ex,xml,re.S)
print(xml_list)
text='\t'.join(xml_list)
print(text)